import os
import aiohttp
import json
import hashlib
import google.generativeai as genai
import numpy as np
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from pydantic_settings import BaseSettings
import uvicorn
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain_text_splitters import TokenTextSplitter
from langchain_community.vectorstores.faiss import FAISS
from langchain_core.documents import Document as LangchainDocument
import tempfile
import asyncio
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor
import faiss
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import shutil
import logging

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Settings
class Settings(BaseSettings):
    gemini_api_key: str = os.getenv("GEMINI_API_KEY", "")
    bearer_token: str = os.getenv("BEARER_TOKEN", "")
    llm_model: str = "gemini-1.5-pro"
    embedding_model: str = "models/embedding-001"
    chunk_size: int = 512
    chunk_overlap: int = 64
    api_host: str = "0.0.0.0"
    api_port: int = int(os.getenv("PORT", 8000))
    debug: bool = os.getenv("ENVIRONMENT", "development") != "production"
    data_dir: str = "./data"

    class Config:
        env_file = ".env"
        extra = "ignore"

settings = Settings()
os.makedirs(settings.data_dir, exist_ok=True)

# Gemini setup
genai.configure(api_key=settings.gemini_api_key)

# Thread pool executor
_executor = ThreadPoolExecutor(max_workers=10)

# Models
class BatchQueryRequest(BaseModel):
    documents: str  # URL to document or email content
    questions: List[str]

class BatchQueryResponse(BaseModel):
    answers: List[str]

# Dummy embedding function for FAISS compatibility
class DummyEmbeddings:
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        raise NotImplementedError("Embeddings are pre-computed")
    def embed_query(self, text: str) -> List[float]:
        raise NotImplementedError("Embeddings are pre-computed")

# Helper functions
def get_document_type(input_str: str) -> str:
    """Determine the document type based on the input string."""
    if input_str.startswith("http") or input_str.startswith("https"):
        path = urlparse(input_str).path
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            return "pdf"
        elif ext == ".docx":
            return "docx"
        else:
            return "unknown"
    else:
        return "email"

def clean_text(text: str) -> str:
    """Clean extracted text by removing extra spaces and normalizing encoding."""
    return " ".join(text.split()).strip()

# Document processing functions
async def process_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """Process PDF files using pdfplumber."""
    chunks = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    chunks.append({
                        "content": clean_text(text),
                        "metadata": {"type": "text", "page": page_num}
                    })
                tables = page.extract_tables()
                for table in tables:
                    chunks.append({
                        "content": json.dumps(table),
                        "metadata": {"type": "table", "page": page_num}
                    })
    except Exception as e:
        logging.error(f"Error processing PDF {pdf_path}: {e}")
    return chunks

async def process_docx(docx_path: str) -> List[Dict[str, Any]]:
    """Process DOCX files using python-docx."""
    chunks = []
    try:
        doc = DocxDocument(docx_path)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        text = "\n".join(full_text)
        text_splitter = TokenTextSplitter(chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap)
        split_texts = text_splitter.split_text(text)
        for split_text in split_texts:
            chunks.append({
                "content": clean_text(split_text),
                "metadata": {"type": "text"}
            })
        # Extract tables
        for table_num, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            chunks.append({
                "content": json.dumps(table_data),
                "metadata": {"type": "table", "table_number": table_num}
            })
    except Exception as e:
        logging.error(f"Error processing DOCX {docx_path}: {e}")
    return chunks

async def process_email(email_content: str) -> List[Dict[str, Any]]:
    """Process email content (text or HTML)."""
    try:
        if "<html" in email_content.lower():
            soup = BeautifulSoup(email_content, "html.parser")
            text = soup.get_text()
        else:
            text = email_content
        text_splitter = TokenTextSplitter(chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap)
        split_texts = text_splitter.split_text(clean_text(text))
        chunks = [{"content": chunk, "metadata": {"type": "text", "source": "email"}} for chunk in split_texts]
    except Exception as e:
        logging.error(f"Error processing email content: {e}")
        return []
    return chunks

class QueryRetrievalSystem:
    def __init__(self):
        self.llm = ChatGoogleGenerativeAI(model=settings.llm_model, google_api_key=settings.gemini_api_key)
        self.prompt_template = PromptTemplate(
            input_variables=["query", "clauses"],
            template="""Based on the provided clauses, answer the query in detail, citing specific parts of the clauses where relevant. Return a JSON object with the following structure:
{{
  "answer": "string",
  "meets_criteria": boolean,
  "applicable_conditions": ["string"],
  "rationale": "string",
  "confidence_score": number,
  "supporting_evidence": ["string"]
}}
Query: {query}

Clauses: {clauses}"""
        )
        self.llm_chain = self.prompt_template | self.llm
        self.vector_store_cache = {}
        self.dummy_embeddings = DummyEmbeddings()

    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        loop = asyncio.get_event_loop()

        def _embed(text: str):
            try:
                result = genai.embed_content(
                    model=settings.embedding_model,
                    content=text,
                    task_type="RETRIEVAL_DOCUMENT"
                )
                return result["embedding"]
            except Exception as e:
                logging.error(f"Embedding failed: {e}")
                return None

        results = await asyncio.gather(*[loop.run_in_executor(_executor, _embed, text) for text in texts])
        valid = [r for r in results if r is not None]
        if not valid:
            raise ValueError("No valid embeddings generated.")
        return valid

    def cosine_similarity(self, a: List[float], b: List[float]) -> float:
        a, b = np.array(a), np.array(b)
        return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

    async def process_document(self, input_str: str) -> FAISS:
        if input_str in self.vector_store_cache:
            return self.vector_store_cache[input_str]

        cache_id = hashlib.sha256(input_str.encode()).hexdigest()
        cache_path = os.path.join(settings.data_dir, cache_id)
        if os.path.exists(cache_path):
            try:
                vs = FAISS.load_local(cache_path, embeddings=self.dummy_embeddings, allow_dangerous_deserialization=True)
                self.vector_store_cache[input_str] = vs
                return vs
            except Exception as e:
                logging.error(f"Failed to load cached vector store for {input_str}: {e}")

        doc_type = get_document_type(input_str)
        if doc_type in ["pdf", "docx"]:
            # Fetch the content
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.get(input_str, timeout=30) as response:
                        response.raise_for_status()
                        content = await response.read()
            except Exception as e:
                logging.error(f"Failed to fetch document {input_str}: {e}")
                raise HTTPException(status_code=400, detail=f"Failed to fetch document: {e}")

            # Save to temporary file
            suffix = ".pdf" if doc_type == "pdf" else ".docx"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            # Process based on type
            try:
                if doc_type == "pdf":
                    chunks = await process_pdf(temp_file_path)
                elif doc_type == "docx":
                    chunks = await process_docx(temp_file_path)
            finally:
                try:
                    os.remove(temp_file_path)
                except Exception as e:
                    logging.warning(f"Failed to delete {temp_file_path}: {e}")
        elif doc_type == "email":
            chunks = await process_email(input_str)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

        if not chunks:
            raise ValueError(f"No content extracted from document: {input_str}")

        chunk_texts = [chunk["content"] for chunk in chunks]
        embeddings = await self.generate_embeddings(chunk_texts)
        if not embeddings or len(embeddings) != len(chunk_texts):
            raise ValueError("Failed to embed all chunks.")

        documents = [
            LangchainDocument(page_content=chunk["content"], metadata={**chunk["metadata"], "embedding": emb})
            for chunk, emb in zip(chunks, embeddings)
        ]

        embeddings_array = np.array(embeddings).astype('float32')
        dimension = embeddings_array.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings_array)

        vector_store = FAISS(
            embedding_function=self.dummy_embeddings,
            index=index,
            docstore={i: documents[i] for i in range(len(documents))},
            index_to_docstore_id={i: i for i in range(len(documents))}
        )

        try:
            vector_store.save_local(cache_path)
            self.vector_store_cache[input_str] = vector_store
        except Exception as e:
            logging.error(f"Failed to save vector store for {input_str}: {e}")
            raise
        return vector_store

    def rerank_chunks(self, query: str, docs: List[LangchainDocument], query_emb: np.ndarray) -> List[LangchainDocument]:
        chunk_embs = [doc.metadata["embedding"] for doc in docs]
        scores = [self.cosine_similarity(query_emb, np.array(emb)) for emb in chunk_embs]
        ranked = sorted(zip(scores, docs), key=lambda x: x[0], reverse=True)
        return [doc for _, doc in ranked[:5]]

    async def process_batch_queries(self, document_url: str, questions: List[str]) -> List[str]:
        try:
            vector_store = await self.process_document(document_url)
            index = vector_store.index
            docstore = vector_store.docstore

            async def handle(q: str) -> str:
                try:
                    query_emb = await self.generate_embeddings([q])
                    if not query_emb:
                        return "Failed to embed query"
                    query_emb = np.array(query_emb[0]).astype('float32').reshape(1, -1)

                    D, I = index.search(query_emb, k=10)
                    docs = [docstore[i] for i in I[0] if i in docstore]

                    reranked = self.rerank_chunks(q, docs, query_emb[0])
                    text = "\n".join([doc.page_content for doc in reranked])

                    try:
                        response = await self.llm_chain.ainvoke({"query": q, "clauses": text})
                        return self.parse_llm_response(response).get("answer", "Unable to determine")
                    except Exception as e:
                        logging.error(f"LLM error for query {q}: {e}")
                        return "LLM error"
                except Exception as e:
                    logging.error(f"Error processing query {q}: {e}")
                    return "Query processing error"

            answers = await asyncio.gather(*(handle(q) for q in questions))
            return answers
        finally:
            self.flush_vector_store(document_url)

    def parse_llm_response(self, response: Any) -> Dict[str, Any]:
        if hasattr(response, "content"):
            response = response.content
        response = str(response).strip()
        if response.startswith("```json"):
            response = response[7:-3].strip()
        elif response.startswith("```"):
            response = response[3:-3].strip()
        try:
            return json.loads(response)
        except json.JSONDecodeError:
            logging.error("[JSON ERROR] Failed to parse JSON from LLM.")
            return {"answer": "Invalid JSON"}

    def flush_vector_store(self, input_str: str):
        """Clear the vector store cache and remove cached files."""
        if input_str in self.vector_store_cache:
            del self.vector_store_cache[input_str]
            logging.info(f"Cleared in-memory vector store cache for {input_str}")
        cache_id = hashlib.sha256(input_str.encode()).hexdigest()
        cache_path = os.path.join(settings.data_dir, cache_id)
        if os.path.exists(cache_path):
            try:
                shutil.rmtree(cache_path)
                logging.info(f"Deleted cached vector store at {cache_path}")
            except Exception as e:
                logging.error(f"Failed to delete cache directory {cache_path}: {e}")

# FastAPI
app = FastAPI(
    title="LLM-Powered Query-Retrieval System",
    description="Query legal/financial docs with LLM",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

# Middleware for Bearer token authentication
@app.middleware("http")
async def authenticate_bearer_token(request: Request, call_next):
    if request.method == "POST" and request.url.path == "/hackrx/run":
        auth_header = request.headers.get("Authorization")
        if not auth_header:
            raise HTTPException(status_code=401, detail="Authorization header missing")
        if not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Invalid Authorization header format")
        token = auth_header[len("Bearer "):].strip()
        if token != settings.bearer_token:
            raise HTTPException(status_code=401, detail="Invalid Bearer token")
    response = await call_next(request)
    return response

query_system = QueryRetrievalSystem()

@app.post("/hackrx/run", response_model=BatchQueryResponse)
async def process_batch_queries(request: BatchQueryRequest):
    try:
        answers = await query_system.process_batch_queries(request.documents, request.questions)
        return BatchQueryResponse(answers=answers)
    except Exception as e:
        logging.error(f"Error in process_batch_queries: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {e}")

@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    logging.error(f"[ERROR] Unhandled Exception: {exc}")
    return JSONResponse(status_code=500, content={"error": "Internal server error", "detail": str(exc)})

if __name__ == "__main__":
    uvicorn.run("app:app", host=settings.api_host, port=settings.api_port, reload=settings.debug)